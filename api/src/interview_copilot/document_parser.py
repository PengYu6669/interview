from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

import pymupdf
from docx import Document


class UnsupportedDocumentError(ValueError):
    pass


class InvalidDocumentError(ValueError):
    pass


class ProtectedDocumentError(ValueError):
    pass


@dataclass(slots=True)
class ParsedDocument:
    filename: str
    media_type: str
    text: str
    page_count: int | None = None
    pages: list[str] | None = None


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()

    if not content:
        raise InvalidDocumentError("文档内容为空")

    if suffix == ".pdf":
        if not content.startswith(b"%PDF-"):
            raise InvalidDocumentError("文件内容不是有效的 PDF")
        try:
            with pymupdf.open(stream=content, filetype="pdf") as document:
                if document.needs_pass:
                    raise ProtectedDocumentError("暂不支持受密码保护的 PDF")
                pages = [page.get_text("text", sort=True).strip() for page in document]
                text = "\n\n".join(page for page in pages if page)
                return ParsedDocument(
                    filename,
                    "application/pdf",
                    text.strip(),
                    len(document),
                    pages,
                )
        except (pymupdf.FileDataError, RuntimeError) as exc:
            raise InvalidDocumentError("PDF 已损坏或无法解析") from exc

    if suffix == ".docx":
        try:
            with ZipFile(BytesIO(content)) as archive:
                entries = set(archive.namelist())
                if "[Content_Types].xml" not in entries or "word/document.xml" not in entries:
                    raise InvalidDocumentError("文件内容不是有效的 DOCX")
            document = Document(BytesIO(content))
        except (BadZipFile, KeyError, ValueError) as exc:
            raise InvalidDocumentError("DOCX 已损坏或无法解析") from exc
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append("\t".join(cell.text.strip() for cell in row.cells))
        return ParsedDocument(
            filename,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "\n".join(paragraphs).strip(),
        )

    if suffix in {".txt", ".md"}:
        try:
            text = content.decode("utf-8-sig").strip()
        except UnicodeDecodeError as exc:
            raise InvalidDocumentError("文本文档必须使用 UTF-8 编码") from exc
        return ParsedDocument(filename, "text/markdown" if suffix == ".md" else "text/plain", text)

    raise UnsupportedDocumentError("仅支持 PDF、DOCX、Markdown 和纯文本文件")
